"""Remove vertical (left/right) cell borders from every table in the
JBI submission DOCX. Elsevier's Guide for Authors mandates "Avoid
vertical rules and shading within table cells".

Keeps top/bottom borders (these are still allowed). Re-saves over the
input file.
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

DOCX = Path(__file__).resolve().parents[2] / "Paper" / "MedFollow_JBI_submission.docx"


def kill_vertical_borders(cell) -> int:
    """Set left and right borders to 'nil' on a cell. Returns number of border attrs touched."""
    tc_pr = cell._tc.find(qn("w:tcPr"))
    if tc_pr is None:
        tc_pr = OxmlElement("w:tcPr")
        cell._tc.insert(0, tc_pr)
    tc_borders = tc_pr.find(qn("w:tcBorders"))
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    n = 0
    for side in ("left", "right"):
        el = tc_borders.find(qn(f"w:{side}"))
        if el is None:
            el = OxmlElement(f"w:{side}")
            tc_borders.append(el)
        if el.get(qn("w:val")) != "nil":
            el.set(qn("w:val"), "nil")
            n += 1
    return n


def kill_table_borders(table) -> int:
    """Also remove insideV (inside vertical) borders at the table level."""
    tbl_pr = table._tbl.find(qn("w:tblPr"))
    if tbl_pr is None:
        return 0
    tbl_borders = tbl_pr.find(qn("w:tblBorders"))
    if tbl_borders is None:
        return 0
    n = 0
    for side in ("left", "right", "insideV"):
        el = tbl_borders.find(qn(f"w:{side}"))
        if el is None:
            el = OxmlElement(f"w:{side}")
            tbl_borders.append(el)
        if el.get(qn("w:val")) != "nil":
            el.set(qn("w:val"), "nil")
            n += 1
    return n


def main() -> None:
    doc = Document(DOCX)
    n_cells = 0
    n_table_attrs = 0
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if kill_vertical_borders(cell):
                    n_cells += 1
        n_table_attrs += kill_table_borders(table)
    doc.save(DOCX)
    print(f"Stripped vertical borders on {n_cells} cell(s) and {n_table_attrs} table-level border attrs across {len(doc.tables)} tables.")


if __name__ == "__main__":
    main()
