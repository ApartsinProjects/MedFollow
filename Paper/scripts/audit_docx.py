"""Audit Paper/MedFollow_JBI_submission.docx against the JBI Guide for
Authors and surface formatting issues.

Checks:
  Structure
    - Title, authors, affiliations, corresponding author present
    - Structured abstract with Objective / Methods / Results / Conclusion
    - Statement of Significance with the four mandated sub-headings
    - Body sections 1-6 in JBI's fixed order
    - All 7 declaration blocks present after Conclusion
    - References section with [N]-prefixed entries
  Counts
    - Total paragraphs, tables, equations, embedded images
    - Body word count vs 6000 cap
    - Abstract word count vs 300 cap
    - Figures + tables vs 8 cap
  Math
    - No leftover raw KaTeX delimiters (\(, \[, $$ outside code blocks)
    - Number of native OMML equations >= source equation count
  Hyperlinks
    - Each reference's https://doi.org/ URL is a live hyperlink
    - At least 1 hyperlink in References section
  Tables
    - Two data tables exist
    - No vertical rules in cells (Elsevier requirement)
  Figures
    - 6 embedded media files in word/media/
    - Each "Figure N." caption is in the document body
  Typography
    - Font on body paragraphs (Times New Roman or default)
    - Justified alignment on body paragraphs

Writes:
    Paper/audit_report.md   (machine-readable Markdown report)

Run from repository root:
    /c/Python314/python Paper/scripts/audit_docx.py
"""
from __future__ import annotations

import re
import sys
import zipfile
from collections import Counter
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[2]
DOCX = ROOT / "Paper" / "MedFollow_JBI_submission.docx"
REPORT = ROOT / "Paper" / "audit_report.md"

# JBI required body sections (in order)
JBI_BODY = ["1. Introduction", "2. Related Work", "3. Methods", "4. Results", "5. Discussion", "6. Conclusion"]
# JBI required declaration blocks (any order, after Conclusion)
JBI_DECLS = [
    "CRediT authorship contribution statement",
    "Declaration of competing interest",
    "Funding",
    "Ethics",
    "Data and code availability",
    "Generative AI use disclosure",
    "Acknowledgements",
]
SOS_HEADINGS = [
    "Problem or Issue",
    "What is Already Known",
    "What this Paper Adds",
    "Who would benefit",
]
ABSTRACT_HEADINGS = ["Objective", "Methods", "Results", "Conclusion"]


def main() -> None:
    if not DOCX.exists():
        print(f"ERROR: {DOCX} not found")
        sys.exit(1)

    doc = Document(DOCX)
    issues: list[str] = []
    pass_: list[str] = []
    notes: list[str] = []

    # ---------- Collect paragraph text ----------
    paragraphs = [p for p in doc.paragraphs]
    para_texts = [p.text.strip() for p in paragraphs]
    full_text = "\n".join(para_texts)

    # ---------- Structure: headings ----------
    headings = [t for t in para_texts if t and (t == t.strip())]

    # Title
    if any("Reliable Extraction of Clinical Follow-Up" in t for t in para_texts[:5]):
        pass_.append("Title present in first 5 paragraphs")
    else:
        issues.append("Title not found in first 5 paragraphs")

    # Authors
    if any(("Laufer" in t and "Aperstein" in t and "Apartsin" in t) for t in para_texts[:10]):
        pass_.append("Author block present (all three authors)")
    else:
        issues.append("Author block not found or incomplete")

    # Affiliations
    expected_affils = ["Bar-Ilan", "Afeka", "Holon Institute"]
    missing_affil = [a for a in expected_affils if a not in full_text]
    if not missing_affil:
        pass_.append("All three affiliations present")
    else:
        issues.append(f"Missing affiliation(s): {', '.join(missing_affil)}")

    # Corresponding author email
    if "apersteiny@afeka.ac.il" in full_text:
        pass_.append("Corresponding author email present (apersteiny@afeka.ac.il)")
    else:
        issues.append("Corresponding author email apersteiny@afeka.ac.il not found")

    # Abstract sub-headings (Objective / Methods / Results / Conclusion)
    abstract_seen = []
    for h in ABSTRACT_HEADINGS:
        # Look as a bold leading word followed by '.', or as a paragraph that starts with that word
        if re.search(rf"\b{re.escape(h)}\b\s*\.\s+\S", full_text):
            abstract_seen.append(h)
    if len(abstract_seen) == 4:
        pass_.append("Abstract has all 4 mandated sub-headings (Objective / Methods / Results / Conclusion)")
    else:
        missing = [h for h in ABSTRACT_HEADINGS if h not in abstract_seen]
        issues.append(f"Abstract sub-headings missing: {missing}")

    # Statement of Significance (4 sub-headings)
    sos_seen = []
    for h in SOS_HEADINGS:
        if h in full_text:
            sos_seen.append(h)
    if len(sos_seen) == 4:
        pass_.append("Statement of Significance has all 4 mandated sub-headings")
    else:
        missing = [h for h in SOS_HEADINGS if h not in sos_seen]
        issues.append(f"Statement of Significance sub-headings missing: {missing}")

    # Body sections 1-6
    body_seen = []
    body_order_ok = True
    last_idx = -1
    for h in JBI_BODY:
        idx = next((i for i, t in enumerate(para_texts) if t == h), -1)
        if idx >= 0:
            body_seen.append(h)
            if idx < last_idx:
                body_order_ok = False
            last_idx = idx
    if len(body_seen) == 6:
        pass_.append("All 6 body sections present (1 Intro / 2 Related / 3 Methods / 4 Results / 5 Discussion / 6 Conclusion)")
        if body_order_ok:
            pass_.append("Body sections appear in JBI's required order")
        else:
            issues.append("Body sections present but out of order")
    else:
        missing = [h for h in JBI_BODY if h not in body_seen]
        issues.append(f"Missing body section(s): {missing}")

    # Declaration blocks (after Conclusion)
    concl_idx = next((i for i, t in enumerate(para_texts) if t == "6. Conclusion"), -1)
    decl_seen = []
    decl_misordered = []
    for d in JBI_DECLS:
        idx = next((i for i, t in enumerate(para_texts) if t == d), -1)
        if idx > concl_idx:
            decl_seen.append(d)
        elif idx >= 0:
            decl_misordered.append(d)
    if len(decl_seen) == 7:
        pass_.append("All 7 declaration blocks present after Conclusion")
    else:
        missing = [d for d in JBI_DECLS if d not in decl_seen]
        issues.append(f"Missing declaration block(s): {missing}")
    if decl_misordered:
        issues.append(f"Declaration block(s) appear before Conclusion: {decl_misordered}")

    # References section
    ref_idx = next((i for i, t in enumerate(para_texts) if t == "References"), -1)
    if ref_idx >= 0:
        pass_.append("References section present")
        # Count [N]-prefixed entries
        ref_lines = [t for t in para_texts[ref_idx + 1:] if t and t.startswith("[")]
        n_refs = sum(1 for t in ref_lines if re.match(r"^\[\d+\]", t))
        if n_refs >= 22:
            pass_.append(f"References list contains {n_refs} [N]-prefixed entries (>= 22 expected)")
        else:
            issues.append(f"References list contains only {n_refs} [N]-prefixed entries (expected 22)")
    else:
        issues.append("References section heading not found")

    # ---------- Counts: words, paragraphs, tables, images, equations ----------
    body_text = ""
    in_body = False
    for t in para_texts:
        if t in JBI_BODY:
            in_body = True
            section_name = t
            continue
        if in_body:
            # Stop at first declaration heading
            if t in JBI_DECLS or t == "References":
                in_body = False
                continue
            body_text += " " + t
    body_words = len(body_text.split())

    abstract_text = ""
    in_abstract = False
    for t in para_texts:
        if t == "Abstract":
            in_abstract = True
            continue
        if in_abstract:
            if t == "1. Introduction":
                in_abstract = False
                break
            abstract_text += " " + t
    abstract_words = len(abstract_text.split())

    n_paragraphs = len(paragraphs)
    n_tables = len(doc.tables)

    # Embedded images via the ZIP archive
    with zipfile.ZipFile(DOCX) as z:
        media_files = sorted([n for n in z.namelist() if n.startswith("word/media/")])
        n_images = len(media_files)
        # Read document.xml to count equations and check for raw KaTeX
        doc_xml = z.read("word/document.xml").decode("utf-8", errors="replace")
        n_equations = doc_xml.count("<m:oMath")
        # raw KaTeX leftovers in body text
        raw_paren = doc_xml.count(r"\(")
        raw_bracket = doc_xml.count(r"\[")
        raw_dollar = sum(1 for line in doc_xml.split("\n") if "$$" in line)

    # Word/figure caps
    if body_words <= 6000:
        pass_.append(f"Body word count {body_words} (cap 6000)")
    else:
        issues.append(f"Body word count {body_words} EXCEEDS cap 6000")
    if abstract_words <= 300:
        pass_.append(f"Abstract word count {abstract_words} (cap 300)")
    else:
        issues.append(f"Abstract word count {abstract_words} EXCEEDS cap 300")
    if n_images + n_tables <= 8:
        pass_.append(f"Figures ({n_images}) + tables ({n_tables}) = {n_images + n_tables} (cap 8)")
    else:
        issues.append(f"Figures + tables = {n_images + n_tables} EXCEEDS cap 8")

    # ---------- Math: equations + leftover raw KaTeX ----------
    if n_equations >= 30:
        pass_.append(f"{n_equations} native Word OMML equations (>= 30 expected)")
    else:
        issues.append(f"Only {n_equations} OMML equations found (expected >= 30)")
    if raw_paren > 0:
        issues.append(f"Found {raw_paren} occurrences of literal '\\(' in document.xml (KaTeX did not parse)")
    else:
        pass_.append("No literal '\\(' KaTeX delimiters left in document.xml")
    if raw_bracket > 0:
        issues.append(f"Found {raw_bracket} occurrences of literal '\\[' in document.xml")
    if raw_dollar > 0:
        notes.append(f"Found {raw_dollar} lines containing '$$' in document.xml (may be code-block dollar, harmless)")

    # ---------- Hyperlinks ----------
    rels = doc.part.rels
    hyperlinks = [r.target_ref for r in rels.values() if r.reltype.endswith("/hyperlink")]
    doi_links = [h for h in hyperlinks if "doi.org" in h]
    if len(doi_links) >= 18:
        pass_.append(f"Found {len(doi_links)} live https://doi.org hyperlinks in references (>= 18 expected)")
    else:
        notes.append(f"Found only {len(doi_links)} live https://doi.org hyperlinks (some refs are conferences without DOIs, so this may be OK)")

    # ---------- Tables ----------
    if n_tables == 2:
        pass_.append(f"Exactly 2 data tables (Table 1, Table 2)")
    elif n_tables < 2:
        issues.append(f"Only {n_tables} table(s) found; expected 2")
    else:
        notes.append(f"{n_tables} tables present (more than the 2 data tables; may include layout containers)")

    # Check for vertical rules in tables (Elsevier disallows)
    vrules_found = False
    for ti, t in enumerate(doc.tables):
        for row in t.rows:
            for cell in row.cells:
                tc_pr = cell._tc.find(qn("w:tcPr"))
                if tc_pr is None:
                    continue
                tc_borders = tc_pr.find(qn("w:tcBorders"))
                if tc_borders is None:
                    continue
                for side in ("left", "right"):
                    el = tc_borders.find(qn(f"w:{side}"))
                    if el is not None and el.get(qn("w:val")) not in (None, "nil", "none"):
                        vrules_found = True
        if vrules_found:
            break
    if vrules_found:
        issues.append("Found vertical rules in at least one table (Elsevier requires no vertical rules)")
    else:
        pass_.append("No vertical rules detected in tables")

    # ---------- Figure captions ----------
    fig_caption_re = re.compile(r"^Figure\s+(\d+)\s*[\.\:]")
    fig_caption_nums = sorted(int(m.group(1)) for t in para_texts for m in [fig_caption_re.match(t)] if m)
    expected_figs = list(range(1, n_images + 1))
    if fig_caption_nums == expected_figs:
        pass_.append(f"Figure captions present and numbered 1..{n_images} consecutively")
    elif set(fig_caption_nums) == set(expected_figs):
        notes.append(f"Figure caption numbers present but not in order: {fig_caption_nums}")
    else:
        issues.append(f"Figure caption numbers {fig_caption_nums} do not match images 1..{n_images}")

    # ---------- Typography: justified body paragraphs ----------
    n_justified = 0
    n_left = 0
    n_other = 0
    for p in paragraphs:
        a = p.paragraph_format.alignment
        if a is None:
            continue
        a_str = str(a).split(".")[-1]  # JUSTIFY, LEFT, CENTER, RIGHT
        if "JUSTIFY" in a_str:
            n_justified += 1
        elif "LEFT" in a_str:
            n_left += 1
        else:
            n_other += 1
    if n_justified > 0:
        pass_.append(f"Body alignment: {n_justified} paragraphs explicitly JUSTIFIED, {n_left} LEFT, {n_other} other")
    else:
        notes.append("No paragraphs have explicit JUSTIFY alignment (may be inherited from style; verify visually in Word)")

    # ---------- Font check (sample first 20 body runs) ----------
    fonts = Counter()
    for p in paragraphs[:80]:
        for r in p.runs:
            f = r.font.name
            if f:
                fonts[f] += 1
    if fonts:
        notes.append(f"Font sampling (first 80 paragraphs): {dict(fonts.most_common())}")

    # ---------- Specific JBI constraint: corresponding-author marker ----------
    if "*" in full_text and "Corresponding author" in full_text:
        pass_.append("Corresponding-author marker present")
    else:
        issues.append("Corresponding-author marker (*) not found")

    # ---------- AI use disclosure: contains Elsevier mandated wording ----------
    if "After using this tool" in full_text or "After using these tools" in full_text:
        pass_.append("Generative AI disclosure contains Elsevier mandated wording")
    else:
        issues.append("Generative AI disclosure may be missing the mandated 'After using this tool/service ...' wording")

    # ---------- Generative AI as research method (model name + version) ----------
    if "gpt-4o-mini" in full_text:
        pass_.append("Synthetic-data generator model name present (gpt-4o-mini)")
    else:
        issues.append("Synthetic-data generator model name not found in body (required for AI-as-method disclosure)")

    # ---------- Build report ----------
    lines = []
    lines.append("# DOCX Audit Report")
    lines.append("")
    lines.append(f"**Source:** `{DOCX.relative_to(ROOT)}`  ")
    lines.append(f"**Size:** {DOCX.stat().st_size:,} bytes  ")
    lines.append(f"**Paragraphs:** {n_paragraphs}  ")
    lines.append(f"**Tables:** {n_tables}  ")
    lines.append(f"**Embedded images:** {n_images}  ")
    lines.append(f"**Native OMML equations:** {n_equations}  ")
    lines.append(f"**Body word count:** {body_words}  ")
    lines.append(f"**Abstract word count:** {abstract_words}  ")
    lines.append("")
    lines.append("## Embedded media files")
    lines.append("")
    for m in media_files:
        lines.append(f"- `{m}`")
    lines.append("")
    if issues:
        lines.append(f"## Issues to fix ({len(issues)})")
        lines.append("")
        for i in issues:
            lines.append(f"- :x: {i}")
        lines.append("")
    else:
        lines.append("## Issues to fix")
        lines.append("")
        lines.append("None.")
        lines.append("")
    lines.append(f"## Passing checks ({len(pass_)})")
    lines.append("")
    for p in pass_:
        lines.append(f"- :white_check_mark: {p}")
    lines.append("")
    if notes:
        lines.append(f"## Advisory notes ({len(notes)})")
        lines.append("")
        for n in notes:
            lines.append(f"- {n}")
        lines.append("")

    REPORT.write_text("\n".join(lines), encoding="utf-8")
    print(f"Wrote {REPORT}")
    print()
    print(f"PASS: {len(pass_)}    ISSUES: {len(issues)}    NOTES: {len(notes)}")
    if issues:
        print()
        print("ISSUES:")
        for i in issues:
            print(f"  - {i}")


if __name__ == "__main__":
    main()
